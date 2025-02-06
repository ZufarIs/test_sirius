import docx
from PIL import Image
import io
import base64
from smolagents import (
    CodeAgent,
    ToolCallingAgent,
    ManagedAgent,
    DuckDuckGoSearchTool,
    VisitWebpageTool,
    HfApiModel,
    Tool,
    tool,
    GradioUI
)
import os

model = HfApiModel(
    "Qwen/Qwen2.5-Coder-32B-Instruct",
    token=os.getenv("HF_API_KEY")
)

image_generation_tool = Tool.from_space(
    "black-forest-labs/FLUX.1-schnell",
    name="image_generator",
    description="Generate an image from a prompt. Returns path to generated image file."
)

@tool
def generate_and_save_image(prompt: str) -> str:
    '''
    Generates an image from a prompt and saves it to the current directory.
    
    Args:
        prompt: prompt to generate the image from
        
    Returns:
        path to the saved image file
    '''
    image_path = image_generation_tool(prompt)
    
    # Move file to working directory
    filename = os.path.basename(image_path)
    new_path = os.path.join(os.getcwd(), filename)
    os.rename(image_path, new_path)
    
    print(f"Image generated and saved to {new_path}")
    return new_path

@tool
def save_to_docx(text: str, filename:str) -> str:
    '''
    Saves the text to a .docx file
    
    Returns a path to the saved .docx file
    
    Args:
        text: text to save to the .docx file
        filename: path to the file to save the text
    '''
    doc = docx.Document()
    doc.add_heading('Пост для соцсетей о пользе раннего развития', 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
        # doc.add_paragraph('Новый абзац')
    doc.save(filename)
    print(f"File saved to {filename}")
    return filename

agent = ToolCallingAgent(
    tools=[DuckDuckGoSearchTool(), VisitWebpageTool(), save_to_docx, generate_and_save_image],
    model=model, 
    add_base_tools=True,
)
managed_scrape_and_write_agent = ManagedAgent(
    agent=agent,
    name="scrape_and_image_gen_agent",
    description="This is an agent that can do web search, scrape web pages, save the text to a .docx file and generate an image based on the text.",
)
manager_agent = CodeAgent(
    tools=[],
    model=model,
    managed_agents=[managed_scrape_and_write_agent],
    add_base_tools=True,
    max_steps=10,
    additional_authorized_imports=['unicodedata', 'os', 'datetime', 'queue', 'statistics', 
                                   'math', 'itertools', 'stat', 'collections', 
                                   'random', 'time', 're', 'PIL'],
)

q1 = "Найди в поиске статьи про то, как раннее развитие влияет на успешность в жизни.\
    Выбери одну статью, которая наиболее интересна и соответствует теме.\
    Создай пост на основе выбранной статьи длиной 1000 - 2000 символов на русском языке для соцсетей и сохрани его текст в формате .docx с помощью инструмента save_to_docx\
    Создай промпт для генерации картинки, используя пост и создай картинку и сохрани его в формате .png в текущей директории.\
    Напиши краткий ответ в формате: \
        Статья: <ссылка на статью>.\
        Пост: <путь к файлу с постом>.\
        Картинка: <путь к картинке>."
        
q2 = "Найди интересную статью про пользу раннего развития для будущего ребенка. \
    Создай интересный пост для соцсетей длиной 1000 - 2000 символов на русском языке и картинку на основе статьи. Сохрани их в формате .docx и .png в локально в текущей директории (не в памяти).\
    Напиши краткий отчет и укажи ссылку на статью, пост и картинку."        
manager_agent.run(
    q2
)
