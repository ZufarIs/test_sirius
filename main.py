import os
from dotenv import load_dotenv
import docx
from scraper import Scraper
from ai_model import AI_Model
from datetime import datetime

load_dotenv()

def save_to_docx(text: str, filename:str = 'post.docx') -> str:
    '''
    Saves the text to a .docx file
    '''
    doc = docx.Document()
    doc.add_heading('Пост для соцсетей о пользе раннего развития', 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = os.path.join('media', f"post_{timestamp}.docx")   
    doc.save(filename)
    return filename

# 1. Формирование поискового запроса, парсинг страниц с новостями и создание их краткого содержания
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
model = AI_Model()
main_theme = 'Последние новости, связанные с целевой аудиторией - мамами. \
    Основной фокус на новостях, касающиеся важности раннего обучения детей арифметике, скорочтению и другим предметам. \
    Под новостями понимаются статьи, в которых приводятся конкретные факты, цифры, статистики, исследования, отзывы экспертов \
    и примеры, которые можно использовать для создания нового поста.'
model.system_prompt = f"Ты профессиональный редактор журнала для мам и умеешь выбирать из множества новостей наиболее интересные - {main_theme}"
model.user_prompt = f"Сформулируй поисковый запрос для поиска свежих новостей по теме: \n{main_theme}\n\nВ качестве ответа верни только сам поисковый запрос"
model.temperature = 0.75
search_query = model.get_answer()
if not search_query:
    print("Не удалось сформулировать поисковый запрос")
    exit()
print(f"Поисковый запрос: {search_query}\n")
scraper = Scraper(search_query)

# 2. Выбор наиболее интересной статьи
model.user_prompt = f"Выбери из множества статей самую интересную и актуальную для мам, на базе которой можно будет создать интересный и актуальный пост для блога. В качестве ответа верни только url статьи: \n{scraper.pages_summary}"
model.temperature = 0.1
selected_url = model.get_answer()
if selected_url not in scraper.pages_summary:
    print(f"Некорректный url: {selected_url}")
    exit()
print(f"Выбранная статья: \n{scraper.pages_content[selected_url]}")

# 3. Создание поста для блога
model.user_prompt = f"Создай пост до 2000 символов для женского блога на основе текста статьи по следующему алгоритму:\n\
    1. Следует отдавать предпочтение подробному разбору одного факта, аспекта или примера, а не обзору глобального вопроса. \
    2. Выбери наиболее интересный и актуальный факт, аспект или пример, упомянутый в статье. \
    3. Объясни почему этот факт, аспект или пример важен для мам и их детей. \
    4. Простым языком, понятным для мам, объясни научные причины этого факта, аспекта или примера. \
    5. Объясни, как именно этот факт или аспект может быть использован для обучения детей, приведи дополнительные примеры. \
    6. Вставляй побольше символов эмодзи, чтобы привлечь внимание читателей.\
    Не используй подзаголовки. Повествование должно плавно переходить от одного факта к другому. Смысловые разделы должны быть разделены переносами строк. \
    Используй обыденную лексику и обязательно объясняй сложные понятия. \
    Текст статьи для написания поста: \n{scraper.pages_content[selected_url]}"
model.temperature = 0.8
post = model.get_answer()
if not post:
    print("Не удалось создать пост")
    exit()
print(f"Пост для блога: \n{post}\n")
filename = save_to_docx(post)
print(f"File saved to {filename}\n\n")

# 4. Создание промпта для генерации изображения
model.user_prompt = f"Создай задание для художника для генерации изображения на основе поста для блога: \n{post}\n\n\nОсобо уточни использование только фенотипических европеоидов как персонажей изображения. В качестве ответа верни только сам промпт для создания изображения"
model.temperature = 0.75
image_prompt = model.get_answer()
if not image_prompt:
    print("Не удалось создать промпт для создания изображения")
    exit()
print(f"Промпт для создания изображения: \n{image_prompt}\n\n")
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
filename = os.path.join('media', f"image_prompt_{timestamp}.txt")   
with open(filename, 'w', encoding='utf-8') as f:
    f.write(image_prompt)

# 5. Генерация изображения
image_path = model.image_generation(image_prompt)
if not image_path:
    print("Не удалось создать изображение")
    exit()
print(f"Изображение сохранено в: \n{image_path}")




